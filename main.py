import webview
import os
import sys
import json
import random
import string
import csv
import io
from pathlib import Path
from datetime import datetime, timedelta

# PDF / Word
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib import colors
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

# ---------------------------------------------------------------------------
# FICTIONAL DATA HELPERS
# ---------------------------------------------------------------------------

FIRST_NAMES = ["Alex","Jordan","Morgan","Taylor","Casey","Riley","Jamie","Avery",
               "Peyton","Quinn","Drew","Reese","Skyler","Blake","Cameron"]
LAST_NAMES  = ["Rivera","Chen","Okafor","Patel","Novak","Larsen","Osei","Fujita",
               "Delacroix","Vasquez","Thornton","Mbeki","Sorensen","Haddad","Winters"]
COMPANY_SUFFIXES = ["LLC","Inc.","Corp.","Group","Solutions","Partners","Services","Technologies"]
COMPANY_WORDS   = ["Apex","Nexus","Crestview","Riverdale","Pinnacle","Lakewood",
                   "Summit","Horizon","Clearwater","Meridian","Ironwood","Bridgepoint"]
STREETS = ["Main St","Oak Ave","Maple Dr","Cedar Blvd","Elm St","Park Rd","Lake Dr","Ridge Rd"]
CITIES  = [
    ("Springfield","IL","62701"),("Riverdale","NY","10471"),("Lakewood","CO","80214"),
    ("Fairview","TX","75069"),("Maplewood","NJ","07040"),("Cedarville","OH","45314"),
    ("Brookhaven","GA","30319"),("Crestwood","MO","63126")
]
STATES = ["AL","AZ","CA","CO","FL","GA","IL","IN","MA","MI",
          "MN","MO","NC","NJ","NV","NY","OH","OR","PA","TX","VA","WA","WI"]

DISCLAIMER_FOOTER = (
    "SAMPLE — For testing, demo, or training use only. "
    "Not valid for legal, financial, medical, regulatory, or identity purposes."
)
WATERMARK = "TEST DOCUMENT"


def rand_name():
    return f"{random.choice(FIRST_NAMES)} {random.choice(LAST_NAMES)}"

def rand_company():
    return f"{random.choice(COMPANY_WORDS)} {random.choice(COMPANY_WORDS)} {random.choice(COMPANY_SUFFIXES)}"

def rand_address():
    num = random.randint(100, 9999)
    street = random.choice(STREETS)
    city, state, zip_ = random.choice(CITIES)
    return f"{num} {street}", f"{city}, {state} {zip_}"

def rand_phone():
    area = random.randint(200, 999)
    mid  = random.randint(200, 999)
    end  = random.randint(1000, 9999)
    return f"({area}) {mid}-{end}"

def rand_email(name: str, company: str):
    local = name.lower().replace(" ", ".")
    domain = company.split()[0].lower()
    return f"{local}@{domain}.example.com"

def rand_date(start_offset=0, end_offset=365):
    base = datetime.today() + timedelta(days=start_offset)
    delta = random.randint(0, end_offset)
    return (base + timedelta(days=delta)).strftime("%B %d, %Y")

def rand_id(prefix="DOC", length=8):
    chars = string.ascii_uppercase + string.digits
    return prefix + "-" + "".join(random.choices(chars, k=length))

def rand_amount(lo=1000, hi=250000):
    return f"${random.randint(lo, hi):,}.00"


# ---------------------------------------------------------------------------
# HTML DOCUMENT BUILDERS (internal intermediary)
# ---------------------------------------------------------------------------

def _html_wrap(title: str, body: str) -> str:
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
  body {{ font-family: 'Segoe UI', Arial, sans-serif; font-size: 13px; color: #1a1a1a;
          max-width: 820px; margin: 40px auto; padding: 0 32px; line-height: 1.6; }}
  h1   {{ font-size: 20px; margin-bottom: 4px; }}
  h2   {{ font-size: 15px; margin: 24px 0 8px; border-bottom: 1px solid #ccc; padding-bottom: 4px; }}
  table {{ width: 100%; border-collapse: collapse; margin: 12px 0; }}
  th, td {{ border: 1px solid #ccc; padding: 7px 10px; text-align: left; font-size: 12px; }}
  th   {{ background: #f0f0f0; font-weight: 600; }}
  .meta {{ color: #555; font-size: 11px; margin-bottom: 24px; }}
  .watermark {{
    position: fixed; top: 46%; left: 50%; transform: translate(-50%,-50%) rotate(-35deg);
    font-size: 72px; color: rgba(0,0,0,0.06); font-weight: 900; pointer-events: none;
    white-space: nowrap; z-index: 0; user-select: none;
  }}
  .footer {{
    margin-top: 48px; border-top: 1px solid #ddd; padding-top: 10px;
    font-size: 10px; color: #888; text-align: center;
  }}
</style>
</head>
<body>
<div class="watermark">{WATERMARK}</div>
{body}
<div class="footer">{DISCLAIMER_FOOTER}</div>
</body>
</html>"""


def _html_to_plain_lines(title: str, body_html: str) -> list[str]:
    """Strip HTML tags to produce plain text lines for PDF/Word."""
    import re
    text = f"{title}\n{'=' * len(title)}\n\n"
    text += DISCLAIMER_FOOTER + "\n\n"
    clean = re.sub(r'<br\s*/?>', '\n', body_html, flags=re.IGNORECASE)
    clean = re.sub(r'</tr>', '\n', clean, flags=re.IGNORECASE)
    clean = re.sub(r'</th>|</td>', '\t', clean, flags=re.IGNORECASE)
    clean = re.sub(r'</p>|</div>|</h[1-6]>', '\n', clean, flags=re.IGNORECASE)
    clean = re.sub(r'<[^>]+>', '', clean)
    import html as html_mod
    clean = html_mod.unescape(clean)
    text += clean
    return [ln.rstrip() for ln in text.splitlines()]


# ---------------------------------------------------------------------------
# EXPORT HELPERS
# ---------------------------------------------------------------------------

def _export_pdf(out_path: str, title: str, lines: list[str]) -> None:
    c = rl_canvas.Canvas(out_path, pagesize=LETTER)
    width, height = LETTER

    def _new_page() -> float:
        # watermark
        c.saveState()
        c.setFont("Helvetica-Bold", 52)
        c.setFillColor(colors.Color(0, 0, 0, 0.05))
        c.translate(width / 2, height / 2)
        c.rotate(45)
        c.drawCentredString(0, 0, WATERMARK)
        c.restoreState()
        # footer
        c.saveState()
        c.setFont("Helvetica", 7)
        c.setFillColor(colors.Color(0.5, 0.5, 0.5))
        c.drawCentredString(width / 2, 24, DISCLAIMER_FOOTER)
        c.restoreState()
        c.setFont("Helvetica", 10)
        return height - 60

    y = _new_page()
    for line in lines:
        if y < 48:
            c.showPage()
            y = _new_page()
        # tab-separated → spaced out
        display = line.replace('\t', '    ')
        c.drawString(48, y, display[:120])
        y -= 14
    c.save()


def _export_docx(out_path: str, title: str, lines: list[str]) -> None:
    doc = DocxDocument()

    # Footer disclaimer
    for section in doc.sections:
        footer_para = section.footer.paragraphs[0]
        footer_para.text = DISCLAIMER_FOOTER
        for run in footer_para.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Watermark notice at top
    wm = doc.add_paragraph()
    wm_run = wm.add_run(f"[ {WATERMARK} ]")
    wm_run.font.size = Pt(9)
    wm_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # Title
    t = doc.add_heading(title, level=1)
    t.runs[0].font.size = Pt(16)

    # Body
    for line in lines[2:]:  # skip title we already added
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if line.startswith("==") or line.startswith("--"):
            continue
        if stripped.isupper() and len(stripped) < 60:
            h = doc.add_heading(stripped, level=2)
            h.runs[0].font.size = Pt(12)
        else:
            p = doc.add_paragraph(stripped.replace('\t', '    '))
            p.runs[0].font.size = Pt(10) if p.runs else None

    doc.save(out_path)


def _write_doc(out_path: str, fmt: str, title: str, body_html: str) -> None:
    if fmt == "html":
        Path(out_path).write_text(_html_wrap(title, body_html), encoding="utf-8")
    elif fmt == "pdf":
        lines = _html_to_plain_lines(title, body_html)
        _export_pdf(out_path, title, lines)
    elif fmt == "docx":
        lines = _html_to_plain_lines(title, body_html)
        _export_docx(out_path, title, lines)
    elif fmt == "txt":
        lines = _html_to_plain_lines(title, body_html)
        Path(out_path).write_text("\n".join(lines), encoding="utf-8")


def _ext(fmt: str) -> str:
    return {"html": ".html", "pdf": ".pdf", "docx": ".docx", "txt": ".txt"}.get(fmt, ".html")


# ---------------------------------------------------------------------------
# DOCUMENT GENERATORS
# ---------------------------------------------------------------------------

def generate_invoice(config: dict) -> dict:
    vendor   = config.get("vendorName") or rand_company()
    buyer    = config.get("buyerName")  or rand_company()
    v_addr1, v_addr2 = rand_address()
    b_addr1, b_addr2 = rand_address()
    inv_num  = rand_id("INV")
    inv_date = rand_date(-30, 0)
    due_date = rand_date(0, 30)

    items = []
    total = 0
    for _ in range(random.randint(2, 6)):
        desc  = random.choice(["Professional Services","Software License","Consulting Hours",
                               "Support & Maintenance","Implementation Services","Training"])
        qty   = random.randint(1, 20)
        rate  = random.randint(50, 500)
        amt   = qty * rate
        total += amt
        items.append((desc, qty, f"${rate:,.2f}", f"${amt:,.2f}"))

    rows = "".join(
        f"<tr><td>{d}</td><td>{q}</td><td>{r}</td><td>{a}</td></tr>"
        for d, q, r, a in items
    )

    body = f"""
    <h1>INVOICE</h1>
    <div class="meta">
      <strong>Invoice #:</strong> {inv_num} &nbsp;|
      <strong>Date:</strong> {inv_date} &nbsp;|
      <strong>Due:</strong> {due_date}
    </div>
    <h2>From</h2>
    <p><strong>{vendor}</strong><br>{v_addr1}<br>{v_addr2}<br>{rand_phone()}</p>
    <h2>Bill To</h2>
    <p><strong>{buyer}</strong><br>{b_addr1}<br>{b_addr2}<br>{rand_phone()}</p>
    <h2>Line Items</h2>
    <table>
      <tr><th>Description</th><th>Qty</th><th>Unit Rate</th><th>Amount</th></tr>
      {rows}
      <tr><td colspan="3"><strong>Total</strong></td><td><strong>${total:,.2f}</strong></td></tr>
    </table>
    """
    return {"base": inv_num, "title": f"Invoice {inv_num}", "html": body}


def generate_contract(config: dict) -> dict:
    buyer    = config.get("buyerName")  or rand_company()
    vendor   = config.get("vendorName") or rand_company()
    b_addr1, b_addr2 = rand_address()
    v_addr1, v_addr2 = rand_address()
    contract_type = config.get("contractType", "MSA")
    doc_id   = rand_id(contract_type)
    eff_date = rand_date(-90, 0)
    exp_date = rand_date(300, 730)

    body = f"""
    <h1>{contract_type} — Master Service Agreement</h1>
    <div class="meta">
      <strong>Doc ID:</strong> {doc_id} &nbsp;|
      <strong>Effective:</strong> {eff_date} &nbsp;|
      <strong>Expires:</strong> {exp_date}
    </div>
    <h2>Parties</h2>
    <p><strong>Client:</strong> {buyer}, {b_addr1}, {b_addr2}</p>
    <p><strong>Vendor:</strong> {vendor}, {v_addr1}, {v_addr2}</p>
    <h2>1. Scope of Services</h2>
    <p>Vendor agrees to provide professional services as mutually agreed in writing.
    Services shall be performed in a professional manner consistent with industry standards.</p>
    <h2>2. Term</h2>
    <p>This Agreement commences on {eff_date} and continues through {exp_date} unless earlier terminated.</p>
    <h2>3. Compensation</h2>
    <p>Client shall pay Vendor {rand_amount()} annually, invoiced quarterly, due within 30 days of receipt.</p>
    <h2>4. Confidentiality</h2>
    <p>Both parties agree to hold in strict confidence any proprietary information exchanged under this Agreement.</p>
    <h2>5. Governing Law</h2>
    <p>This Agreement shall be governed by the laws of the State of {random.choice(STATES)}.</p>
    <h2>Signatures</h2>
    <table>
      <tr><th>Client</th><th>Vendor</th></tr>
      <tr><td>{rand_name()}<br>{buyer}</td><td>{rand_name()}<br>{vendor}</td></tr>
    </table>
    """
    return {"base": f"{doc_id}_contract", "title": f"{contract_type} {doc_id}", "html": body}


def generate_intake_form(config: dict) -> dict:
    org   = config.get("buyerName") or rand_company()
    doc_id = rand_id("FORM")
    name   = rand_name()
    dob    = rand_date(-365*40, -365*18)
    addr1, addr2 = rand_address()

    body = f"""
    <h1>Intake / Registration Form</h1>
    <div class="meta"><strong>Form ID:</strong> {doc_id} &nbsp;| <strong>Organization:</strong> {org}</div>
    <h2>Applicant Information</h2>
    <table>
      <tr><th>Field</th><th>Value</th></tr>
      <tr><td>Full Name</td><td>{name}</td></tr>
      <tr><td>Date of Birth</td><td>{dob}</td></tr>
      <tr><td>Address</td><td>{addr1}, {addr2}</td></tr>
      <tr><td>Phone</td><td>{rand_phone()}</td></tr>
      <tr><td>Email</td><td>{rand_email(name, org)}</td></tr>
      <tr><td>Preferred Contact</td><td>{random.choice(["Phone","Email","Mail"])}</td></tr>
    </table>
    <h2>Service Request</h2>
    <table>
      <tr><th>Field</th><th>Value</th></tr>
      <tr><td>Service Type</td><td>{random.choice(["New Enrollment","Transfer","Update","Renewal","Inquiry"])}</td></tr>
      <tr><td>Requested Date</td><td>{rand_date(0, 30)}</td></tr>
      <tr><td>Notes</td><td>N/A</td></tr>
    </table>
    <h2>Acknowledgment</h2>
    <p>I certify the information above is accurate to the best of my knowledge.</p>
    <p>Signature: ______________________________ &nbsp; Date: ___________</p>
    """
    return {"base": f"{doc_id}_intake", "title": f"Intake Form {doc_id}", "html": body}


def generate_sop(config: dict) -> dict:
    org    = config.get("buyerName") or rand_company()
    doc_id = rand_id("SOP")
    topic  = random.choice(["Vendor Onboarding","Data Entry Review","Document Archiving",
                            "Access Request","Incident Response","Quality Review"])
    body = f"""
    <h1>Standard Operating Procedure</h1>
    <div class="meta">
      <strong>Doc ID:</strong> {doc_id} &nbsp;|
      <strong>Topic:</strong> {topic} &nbsp;|
      <strong>Organization:</strong> {org}
    </div>
    <h2>Purpose</h2>
    <p>This SOP defines the standard process for {topic.lower()} within {org}.</p>
    <h2>Scope</h2>
    <p>Applies to all staff involved in {topic.lower()} activities.</p>
    <h2>Procedure</h2>
    <table>
      <tr><th>Step</th><th>Action</th><th>Responsible</th></tr>
      {''.join(f"<tr><td>{i+1}</td><td>{a}</td><td>{random.choice(['Manager','Analyst','Admin','Lead'])}</td></tr>"
        for i,a in enumerate([
          "Initiate request and document in tracking system",
          "Review for completeness and accuracy",
          "Escalate to supervisor if exceptions found",
          "Obtain required approvals",
          "Complete processing and file documentation",
          "Confirm completion with requesting party"
        ]))}
    </table>
    <h2>Review Schedule</h2>
    <p>This SOP shall be reviewed annually. Last reviewed: {rand_date(-365, 0)}.</p>
    """
    return {"base": f"{doc_id}_sop", "title": f"SOP {doc_id}", "html": body}


def generate_offer_letter(config: dict) -> dict:
    org    = config.get("buyerName") or rand_company()
    doc_id = rand_id("OFR")
    name   = rand_name()
    role   = random.choice(["Data Analyst","Project Manager","Software Engineer",
                            "Operations Coordinator","Business Analyst","IT Specialist"])
    salary = random.randint(45, 135) * 1000
    start  = rand_date(7, 45)
    body = f"""
    <h1>Employment Offer Letter</h1>
    <div class="meta"><strong>Ref:</strong> {doc_id} &nbsp;| <strong>Date:</strong> {rand_date(-5,0)}</div>
    <p>Dear {name},</p>
    <p>We are pleased to offer you the position of <strong>{role}</strong> at <strong>{org}</strong>.
    This is a full-time position with a start date of <strong>{start}</strong>.</p>
    <h2>Compensation & Benefits</h2>
    <table>
      <tr><th>Item</th><th>Detail</th></tr>
      <tr><td>Annual Salary</td><td>${salary:,}.00</td></tr>
      <tr><td>Benefits</td><td>Medical, Dental, Vision — effective 30 days from start</td></tr>
      <tr><td>PTO</td><td>{random.randint(10,20)} days per year</td></tr>
      <tr><td>Work Location</td><td>{random.choice(["On-site","Remote","Hybrid"])}</td></tr>
    </table>
    <p>Please sign and return this letter by <strong>{rand_date(3,10)}</strong> to confirm your acceptance.</p>
    <p>Sincerely,<br><strong>{rand_name()}</strong><br>Human Resources, {org}</p>
    <p style="margin-top:32px">Accepted: ______________________________ &nbsp; Date: ___________</p>
    """
    return {"base": f"{doc_id}_offer", "title": f"Offer Letter {doc_id}", "html": body}


def generate_purchase_order(config: dict) -> dict:
    buyer  = config.get("buyerName")  or rand_company()
    vendor = config.get("vendorName") or rand_company()
    doc_id = rand_id("PO")
    po_date = rand_date(-10, 0)

    items = []
    total = 0
    for _ in range(random.randint(3, 8)):
        desc = random.choice(["Office Supplies","Hardware","Software License",
                              "Furniture","IT Equipment","Maintenance Services","Subscription"])
        qty  = random.randint(1, 50)
        unit = random.randint(10, 2000)
        amt  = qty * unit
        total += amt
        items.append((desc, qty, f"${unit:,.2f}", f"${amt:,.2f}"))

    rows = "".join(
        f"<tr><td>{d}</td><td>{q}</td><td>{r}</td><td>{a}</td></tr>"
        for d, q, r, a in items
    )
    body = f"""
    <h1>Purchase Order</h1>
    <div class="meta">
      <strong>PO #:</strong> {doc_id} &nbsp;|
      <strong>Date:</strong> {po_date}
    </div>
    <h2>Buyer</h2>
    <p><strong>{buyer}</strong></p>
    <h2>Vendor</h2>
    <p><strong>{vendor}</strong></p>
    <h2>Order Items</h2>
    <table>
      <tr><th>Description</th><th>Qty</th><th>Unit Price</th><th>Total</th></tr>
      {rows}
      <tr><td colspan="3"><strong>Order Total</strong></td><td><strong>${total:,.2f}</strong></td></tr>
    </table>
    <p>Authorized by: ______________________________ &nbsp; Date: ___________</p>
    """
    return {"base": f"{doc_id}_po", "title": f"PO {doc_id}", "html": body}


DOC_GENERATORS = {
    "invoice":        generate_invoice,
    "contract":       generate_contract,
    "intake_form":    generate_intake_form,
    "sop":            generate_sop,
    "offer_letter":   generate_offer_letter,
    "purchase_order": generate_purchase_order,
}


# ---------------------------------------------------------------------------
# DATA FILE GENERATORS
# ---------------------------------------------------------------------------

def _make_header_style(wb):
    return {
        "font": Font(bold=True, color="FFFFFF", size=11),
        "fill": PatternFill("solid", fgColor="2E5D9E"),
        "alignment": Alignment(horizontal="center", wrap_text=True),
        "border": Border(
            bottom=Side(style="thin", color="FFFFFF"),
            right=Side(style="thin", color="FFFFFF"),
        ),
    }

def _apply_header(ws, headers):
    style = _make_header_style(None)
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = style["font"]
        cell.fill = style["fill"]
        cell.alignment = style["alignment"]

def _disclaimer_row(ws, col_count):
    ws.insert_rows(1)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=col_count)
    cell = ws.cell(row=1, column=1, value=DISCLAIMER_FOOTER)
    cell.font = Font(bold=True, color="7A4F1A", size=9)
    cell.fill = PatternFill("solid", fgColor="FEF6EA")
    cell.alignment = Alignment(wrap_text=True)
    ws.row_dimensions[1].height = 22


def _gen_customer_rows(n):
    headers = ["customer_id","first_name","last_name","email","phone",
               "company","address","city","state","zip","account_type","created_date"]
    rows = []
    for i in range(n):
        fn = random.choice(FIRST_NAMES)
        ln = random.choice(LAST_NAMES)
        co = rand_company()
        city, state, zip_ = random.choice(CITIES)
        rows.append([
            f"CUST-{10000+i}", fn, ln, rand_email(f"{fn} {ln}", co),
            rand_phone(), co, f"{random.randint(100,9999)} {random.choice(STREETS)}",
            city, state, zip_,
            random.choice(["Standard","Premium","Enterprise","Trial"]),
            rand_date(-730, 0),
        ])
    return headers, rows

def _gen_vendor_rows(n):
    headers = ["vendor_id","vendor_name","contact_name","email","phone",
               "address","city","state","zip","category","status","contract_exp"]
    rows = []
    for i in range(n):
        co = rand_company()
        cn = rand_name()
        city, state, zip_ = random.choice(CITIES)
        rows.append([
            f"VEN-{2000+i}", co, cn, rand_email(cn, co), rand_phone(),
            f"{random.randint(100,9999)} {random.choice(STREETS)}",
            city, state, zip_,
            random.choice(["Software","Hardware","Consulting","Staffing","Facilities","Legal"]),
            random.choice(["Active","Inactive","Pending","Under Review"]),
            rand_date(0, 730),
        ])
    return headers, rows

def _gen_transaction_rows(n):
    headers = ["transaction_id","date","amount","currency","type","status",
               "account_id","description","reference","processed_by"]
    rows = []
    for i in range(n):
        rows.append([
            f"TXN-{100000+i}",
            rand_date(-365, 0),
            round(random.uniform(10, 50000), 2),
            "USD",
            random.choice(["Payment","Refund","Adjustment","Credit","Debit","Transfer"]),
            random.choice(["Completed","Pending","Failed","Reversed","On Hold"]),
            f"ACC-{random.randint(1000,9999)}",
            random.choice(["Monthly fee","Service charge","Reimbursement",
                           "License renewal","One-time purchase","Wire transfer"]),
            rand_id("REF", 6),
            rand_name(),
        ])
    return headers, rows

def _gen_employee_rows(n):
    headers = ["employee_id","first_name","last_name","email","department",
               "title","hire_date","salary","status","manager","location"]
    rows = []
    depts = ["Engineering","Finance","HR","Operations","Sales","Marketing","Legal","IT","Product"]
    titles = ["Analyst","Senior Analyst","Manager","Director","Coordinator","Specialist","Lead","Associate"]
    for i in range(n):
        fn = random.choice(FIRST_NAMES)
        ln = random.choice(LAST_NAMES)
        co = rand_company()
        dept = random.choice(depts)
        rows.append([
            f"EMP-{5000+i}", fn, ln, rand_email(f"{fn} {ln}", co),
            dept, f"{random.choice(titles)}",
            rand_date(-1825, -30),
            random.randint(42000, 185000),
            random.choice(["Active","On Leave","Terminated","Contractor"]),
            rand_name(),
            random.choice(["Remote","New York, NY","Chicago, IL","Austin, TX","Atlanta, GA"]),
        ])
    return headers, rows

def _gen_inventory_rows(n):
    headers = ["sku","product_name","category","quantity_on_hand","reorder_point",
               "unit_cost","unit_price","supplier","warehouse","last_updated"]
    cats = ["Electronics","Office","Furniture","Software","Supplies","Hardware","Peripherals"]
    products = ["Laptop","Monitor","Keyboard","Mouse","Desk","Chair","Headset",
                "Webcam","Router","Switch","Cable","License","Printer","Scanner"]
    rows = []
    for i in range(n):
        cost = round(random.uniform(5, 2000), 2)
        rows.append([
            f"SKU-{random.randint(10000,99999)}",
            f"{random.choice(products)} {random.choice(['Pro','Plus','Lite','Standard','HD','v2'])}",
            random.choice(cats),
            random.randint(0, 500),
            random.randint(5, 50),
            cost,
            round(cost * random.uniform(1.2, 2.5), 2),
            rand_company(),
            random.choice(["WH-East","WH-West","WH-Central","WH-Remote"]),
            rand_date(-30, 0),
        ])
    return headers, rows

def _gen_messy_rows(n):
    headers = ["id","Name","name ","EMAIL","Phone Number","phone_number",
               "Company","COMPANY","Amount","amount ","Date","DATE","Status","status"]
    rows = []
    for i in range(n):
        name = rand_name()
        co   = rand_company()
        amt  = round(random.uniform(10, 9999), 2)
        dt   = rand_date(-365, 0)
        status = random.choice(["Active","active","ACTIVE","","N/A",None,"Yes","TRUE","1"])
        rows.append([
            random.choice([f"ID-{i}", str(i), f"{i:05d}", None, i]),
            name if random.random() > 0.1 else name.upper(),
            name.lower() if random.random() > 0.2 else None,
            rand_email(name, co) if random.random() > 0.15 else name.split()[0]+"@",
            rand_phone() if random.random() > 0.1 else str(random.randint(1000000000,9999999999)),
            rand_phone() if random.random() > 0.2 else None,
            co, co.upper() if random.random() > 0.3 else co + " ",
            f"${amt:,.2f}", str(amt),
            dt, dt.replace(",", "").replace(" ", "/"),
            status, status,
        ])
    return headers, rows


DATA_GENERATORS = {
    "customers":    _gen_customer_rows,
    "vendors":      _gen_vendor_rows,
    "transactions": _gen_transaction_rows,
    "employees":    _gen_employee_rows,
    "inventory":    _gen_inventory_rows,
    "messy":        _gen_messy_rows,
}


def _write_data(out_dir: str, data_key: str, fmt: str, n: int, config: dict) -> list[str]:
    gen_fn = DATA_GENERATORS.get(data_key)
    if not gen_fn:
        return []
    headers, rows = gen_fn(n)
    created = []

    base = f"{data_key}_{rand_id('', 4).lstrip('-')}"
    disclaimer_row = [DISCLAIMER_FOOTER] + [""] * (len(headers) - 1)

    if fmt in ("csv", "both_data"):
        path = os.path.join(out_dir, base + ".csv")
        with open(path, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(disclaimer_row)
            w.writerow(headers)
            w.writerows(rows)
        created.append(path)

    if fmt in ("xlsx", "both_data"):
        path = os.path.join(out_dir, base + ".xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = data_key.replace("_", " ").title()
        _apply_header(ws, headers)
        alt_fill = PatternFill("solid", fgColor="F0F4FB")
        for r_idx, row in enumerate(rows, 2):
            for c_idx, val in enumerate(row, 1):
                cell = ws.cell(row=r_idx, column=c_idx, value=val)
                if r_idx % 2 == 0:
                    cell.fill = alt_fill
        for col in ws.columns:
            max_len = max((len(str(c.value or "")) for c in col), default=8)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)
        _disclaimer_row(ws, len(headers))
        wb.save(path)
        created.append(path)

    if fmt in ("json", "both_data"):
        path = os.path.join(out_dir, base + ".json")
        data_out = {
            "_disclaimer": DISCLAIMER_FOOTER,
            "generated": datetime.today().isoformat(),
            "records": [dict(zip(headers, row)) for row in rows],
        }
        Path(path).write_text(json.dumps(data_out, indent=2, default=str), encoding="utf-8")
        created.append(path)

    if fmt in ("txt", "both_data"):
        path = os.path.join(out_dir, base + ".txt")
        lines = [DISCLAIMER_FOOTER, ""]
        col_widths = [max(len(str(h)), max((len(str(r[i] or "")) for r in rows), default=0))
                      for i, h in enumerate(headers)]
        sep = "+" + "+".join("-" * (w + 2) for w in col_widths) + "+"
        def fmt_row(vals):
            return "|" + "|".join(f" {str(v or ''):<{col_widths[i]}} " for i, v in enumerate(vals)) + "|"
        lines += [sep, fmt_row(headers), sep]
        for row in rows:
            lines.append(fmt_row(row))
        lines.append(sep)
        Path(path).write_text("\n".join(lines), encoding="utf-8")
        created.append(path)

    return created


# ---------------------------------------------------------------------------
# PYWEBVIEW API
# ---------------------------------------------------------------------------

class API:
    def generate(self, payload: dict):
        output_folder = Path(payload.get("outputFolder") or Path.home() / "Documents" / "sample_docs")
        output_folder.mkdir(parents=True, exist_ok=True)

        doc_types  = payload.get("docTypes", [])
        quantity   = int(payload.get("quantity", 1))
        fmt        = payload.get("format", "pdf").lower()
        mode       = payload.get("mode", "documents")
        results    = []

        # Separate doc vs data keys
        doc_keys  = [k for k in doc_types if k in DOC_GENERATORS]
        data_keys = [k for k in doc_types if k in DATA_GENERATORS]

        # Document export
        doc_fmt = fmt if fmt in ("html", "pdf", "docx", "txt") else "pdf"
        for doc_key in doc_keys:
            gen_fn = DOC_GENERATORS[doc_key]
            for _ in range(quantity):
                result = gen_fn(payload)
                filename = result["base"] + _ext(doc_fmt)
                out_path = str(output_folder / filename)
                _write_doc(out_path, doc_fmt, result["title"], result["html"])
                results.append(out_path)

        # Data export
        data_fmt = fmt if fmt in ("csv", "xlsx", "json", "txt") else "xlsx"
        for data_key in data_keys:
            for _ in range(quantity):
                paths = _write_data(str(output_folder), data_key, data_fmt, 50, payload)
                results.extend(paths)

        # Always also write an Excel summary if docs were generated
        if doc_keys and doc_fmt != "xlsx":
            summary_path = str(output_folder / f"summary_{rand_id('',4).lstrip('-')}.xlsx")
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Generated Files"
            headers = ["#", "File", "Type", "Format", "Generated At"]
            _apply_header(ws, headers)
            for i, fp in enumerate(results, 1):
                fname = os.path.basename(fp)
                ftype = fname.split("_")[0].upper()
                ws.cell(row=i+1, column=1, value=i)
                ws.cell(row=i+1, column=2, value=fname)
                ws.cell(row=i+1, column=3, value=ftype)
                ws.cell(row=i+1, column=4, value=doc_fmt.upper())
                ws.cell(row=i+1, column=5, value=datetime.today().strftime("%Y-%m-%d %H:%M"))
            for col in ws.columns:
                ws.column_dimensions[col[0].column_letter].width = max(
                    len(str(col[0].value or "")),
                    max((len(str(c.value or "")) for c in col[1:]), default=0)
                ) + 4
            _disclaimer_row(ws, len(headers))
            wb.save(summary_path)
            results.append(summary_path)

        return {"success": True, "files": results, "folder": str(output_folder)}

    def choose_folder(self):
        dirs = webview.windows[0].create_file_dialog(webview.FOLDER_DIALOG)
        if dirs and len(dirs) > 0:
            return dirs[0]
        return None

    def open_folder(self, folder_path: str):
        os.startfile(folder_path) if sys.platform == "win32" else os.system(f'open "{folder_path}"')


# ---------------------------------------------------------------------------
# ENTRY POINT
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    api = API()
    ui_path = Path(__file__).parent / "ui.html"
    webview.create_window(
        title="Test Data & Document Generator",
        url=str(ui_path),
        js_api=api,
        width=860,
        height=780,
        min_size=(700, 600),
        resizable=True,
    )
    webview.start(debug=False)
