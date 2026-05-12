import os
import json
import random
import secrets
from dataclasses import dataclass, asdict
from datetime import date, timedelta

import customtkinter as ctk
from tkinter import filedialog, messagebox

from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt

# =========================
# LEGAL / SAFETY CONSTANTS
# =========================

APP_DISCLAIMER = (
    "TEST DOCUMENTS ONLY — Generated content is for testing, demo, or training use only. "
    "Not valid for legal, financial, medical, regulatory, or identity purposes."
)

WATERMARK_TEXT = "TEST DOCUMENT"

FOOTER_NOTICE = (
    "TEST DOCUMENT — For testing, demo, or training use only. "
    "Not valid for legal, financial, medical, regulatory, or identity purposes."
)

# =========================
# DATA MODELS
# =========================

@dataclass
class Party:
    name: str
    address1: str
    address2: str
    phone: str


@dataclass
class SessionData:
    session_id: str
    created_on: str
    buyer: Party
    vendor: Party
    contract_number: str
    effective_date: str
    term_months: int
    annual_fee: int
    governing_law: str
    contract_type: str
    vertical: str


# =========================
# CONSTANTS / RNG
# =========================

CONTRACT_TYPES = ["MSA", "PSA", "BAA", "DUA", "SLA"]
VERTICALS = ["Healthcare", "Higher Education", "Government", "Commercial"]
EXPORT_FORMATS = ["PDF", "Word"]

RNG = random.Random(secrets.randbits(64))

# =========================
# RANDOM HELPERS
# =========================

FIRST_WORDS = [
    "Apex", "Horizon", "Meridian", "Nexus", "Pinnacle",
    "Summit", "Catalyst", "Elevate", "Vantage", "Clarity",
    "Beacon", "Keystone", "Ascend", "Luminary", "Strata",
]

INDUSTRY_WORDS = [
    "Clinical", "Health", "Analytics", "Solutions", "Systems",
    "Informatics", "Care", "Data", "Consulting", "Technologies",
    "Outcomes", "Insights", "Services", "Partners", "Strategies",
]

CORP_SUFFIXES = ["LLC", "Inc.", "Group", "Partners", "Associates"]

STREETS = ["Oak", "Maple", "Commerce", "Innovation", "Enterprise", "Corporate", "Tech", "Riverside", "Lakeside"]
STREET_TYPES = ["Blvd", "Dr", "Ave", "Way", "Pkwy", "St", "Ct"]
CITIES = [
    ("Austin", "TX"), ("Nashville", "TN"), ("Atlanta", "GA"),
    ("Denver", "CO"), ("Charlotte", "NC"), ("Phoenix", "AZ"),
    ("Raleigh", "NC"), ("Tampa", "FL"), ("Columbus", "OH"),
    ("Indianapolis", "IN"),
]


def rand_phone() -> str:
    return f"({RNG.randint(200, 989)}) {RNG.randint(200, 989)}-{RNG.randint(1000, 9999)}"


def rand_address() -> tuple[str, str]:
    city, state = RNG.choice(CITIES)
    street = f"{RNG.randint(100, 9999)} {RNG.choice(STREETS)} {RNG.choice(STREET_TYPES)}"
    return street, f"{city}, {state} {RNG.randint(10000, 99999)}"


def rand_company() -> str:
    return f"{RNG.choice(FIRST_WORDS)} {RNG.choice(INDUSTRY_WORDS)} {RNG.choice(CORP_SUFFIXES)}"


def rand_contract_number() -> str:
    return f"AGR-{RNG.randint(2024, 2026)}-{RNG.randint(10000, 99999)}"


def rand_date_near() -> str:
    d = date.today() + timedelta(days=RNG.randint(-60, 60))
    return d.strftime("%B %d, %Y")


def sanitize_filename(s: str) -> str:
    s = "".join(ch if ch.isalnum() or ch in "._-" else "_" for ch in s.strip())
    return s[:80]


# =========================
# DOCUMENT BUILDERS
# =========================


def build_contract(sd: SessionData, exhibits: list[str]) -> str:
    base = f"""
TEST DOCUMENT — CONTRACT SAMPLE (NOT LEGAL ADVICE)

Contract Number: {sd.contract_number}
Effective Date: {sd.effective_date}

Buyer:
{sd.buyer.name}
{sd.buyer.address1}
{sd.buyer.address2}
{sd.buyer.phone}

Vendor:
{sd.vendor.name}
{sd.vendor.address1}
{sd.vendor.address2}
{sd.vendor.phone}

Term: {sd.term_months} months
Annual Fee: ${sd.annual_fee:,}
Governing Law: {sd.governing_law}

This is a test {sd.contract_type} document for the {sd.vertical} vertical.
It is for testing, demo, and training purposes only.
""".rstrip()

    parts = [base]

    if "A" in exhibits:
        parts.append("""
EXHIBIT A — SCOPE OF SERVICES

Vendor will provide implementation, configuration, and support services as mutually agreed.
""".strip())

    if "B" in exhibits:
        parts.append(f"""
EXHIBIT B — PRICING

Annual Fee: ${sd.annual_fee:,} USD.
Additional services may be billed separately.
""".strip())

    if "C" in exhibits:
        parts.append("""
EXHIBIT C — SECURITY

Vendor will maintain reasonable administrative, technical, and physical safeguards.
""".strip())

    return "\n\n".join(parts)


def build_coi(sd: SessionData) -> str:
    return f"""
TEST DOCUMENT — CERTIFICATE OF INSURANCE

Insured: {sd.vendor.name}
Certificate Holder: {sd.buyer.name}

General Liability: $1,000,000
Professional Liability: $1,000,000

Effective: {sd.effective_date}
""".strip()


def build_amendment(sd: SessionData) -> str:
    return f"""
TEST DOCUMENT — AMENDMENT

Amendment to Contract {sd.contract_number}

Buyer: {sd.buyer.name}
Vendor: {sd.vendor.name}

This amendment extends the agreement by 12 months and may adjust fees.
""".strip()


# =========================
# EXPORTERS
# =========================


def draw_pdf_watermark(c: canvas.Canvas, width: float, height: float) -> None:
    c.saveState()
    c.setFont("Helvetica-Bold", 42)
    c.setFillGray(0.88)
    c.translate(width / 2, height / 2)
    c.rotate(45)
    c.drawCentredString(0, 0, WATERMARK_TEXT)
    c.restoreState()


def draw_pdf_footer(c: canvas.Canvas, width: float) -> None:
    c.saveState()
    c.setFont("Helvetica", 8)
    c.setFillGray(0.35)
    c.drawCentredString(width / 2, 28, FOOTER_NOTICE)
    c.restoreState()


def export_pdf(path: str, title: str, body: str) -> None:
    c = canvas.Canvas(path, pagesize=LETTER)
    width, height = LETTER

    def start_page() -> float:
        draw_pdf_watermark(c, width, height)
        draw_pdf_footer(c, width)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(72, height - 72, title)
        c.setFont("Helvetica", 11)
        return height - 110

    y = start_page()

    for line in body.split("\n"):
        if y < 60:
            c.showPage()
            y = start_page()
        c.drawString(72, y, line)
        y -= 14

    c.save()


def export_docx(path: str, title: str, body: str) -> None:
    doc = Document()

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.text = FOOTER_NOTICE

    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(16)

    notice = doc.add_paragraph()
    notice_run = notice.add_run(WATERMARK_TEXT)
    notice_run.bold = True
    notice_run.font.size = Pt(12)

    doc.add_paragraph("")
    for line in body.split("\n"):
        doc.add_paragraph(line)

    doc.save(path)


# =========================
# GENERATION ENGINE
# =========================


def generate_docs(
    out_dir: str,
    buyer: Party,
    vendor: Party,
    contract_type: str,
    vertical: str,
    fmt: str,
    make_contract: bool,
    make_coi: bool,
    make_amend: bool,
    exhibits: list[str],
) -> list[str]:
    os.makedirs(out_dir, exist_ok=True)

    sd = SessionData(
        session_id=secrets.token_hex(4).upper(),
        created_on=date.today().isoformat(),
        buyer=buyer,
        vendor=vendor,
        contract_number=rand_contract_number(),
        effective_date=rand_date_near(),
        term_months=12,
        annual_fee=RNG.randint(50000, 150000),
        governing_law="New Jersey",
        contract_type=contract_type,
        vertical=vertical,
    )

    created: list[str] = []

    def write(base_name: str, title: str, body: str) -> None:
        base_name = sanitize_filename(base_name)
        ext = ".pdf" if fmt == "PDF" else ".docx"
        path = os.path.join(out_dir, base_name + ext)
        exporter = export_pdf if fmt == "PDF" else export_docx
        exporter(path, title, body)
        created.append(path)

    buyer_slug = sanitize_filename(buyer.name or "Buyer")
    vendor_slug = sanitize_filename(vendor.name or "Vendor")

    if make_contract:
        write(f"{vendor_slug}_to_{buyer_slug}_Contract", "Test Document - Contract", build_contract(sd, exhibits))
    if make_coi:
        write(f"{vendor_slug}_COI", "Test Document - COI", build_coi(sd))
    if make_amend:
        write(f"{vendor_slug}_to_{buyer_slug}_Amendment", "Test Document - Amendment", build_amendment(sd))

    metadata = {
        "document_type": "test",
        "for_testing_only": True,
        "not_for_production_use": True,
        "watermark": WATERMARK_TEXT,
        "footer_notice": FOOTER_NOTICE,
        "generated_by": "Random Document and Data Generator",
        "generated_on": date.today().isoformat(),
        "session_data": asdict(sd),
    }

    meta_path = os.path.join(out_dir, "metadata.json")
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=2)
    created.append(meta_path)

    return created


# =========================
# GUI
# =========================


class App(ctk.CTk):
    def __init__(self):
        super().__init__()

        ctk.set_appearance_mode("light")
        self.ACCENT = "#2AB3A3"
        self.ACCENT_HOVER = "#229C8F"
        self.PANEL_BG = "#FFFFFF"
        self.PANEL_BORDER = "#DDDDDD"

        self.title("Document Generator — Test Documents")
        self.geometry("1100x760")
        self.minsize(1000, 700)

        self.out_dir = os.path.join(os.getcwd(), "sample_docs")

        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)
        for row in range(7):
            self.grid_rowconfigure(row, weight=0)
        self.grid_rowconfigure(0, weight=1)

        self.buyer_frame = self._panel("Your Organization (Buyer / Customer)")
        self.buyer_frame.grid(row=0, column=0, padx=20, pady=20, sticky="nsew")
        self.buyer_frame.grid_columnconfigure(0, weight=1)

        self.buyer_name = self._entry_row(self.buyer_frame, 0, "Company Name (blank = auto)")
        self.buyer_addr1 = self._entry_row(self.buyer_frame, 1, "Address Line 1 (blank = auto)")
        self.buyer_addr2 = self._entry_row(self.buyer_frame, 2, "City, State ZIP (blank = auto)")
        self.buyer_phone = self._entry_row(self.buyer_frame, 3, "Phone (blank = auto)")

        self.vendor_frame = self._panel("Other Organization (Vendor / Service Provider)")
        self.vendor_frame.grid(row=0, column=1, padx=20, pady=20, sticky="nsew")
        self.vendor_frame.grid_columnconfigure(0, weight=1)

        self.vendor_name = self._entry_row(self.vendor_frame, 0, "Company Name (blank = auto)")
        self.vendor_addr1 = self._entry_row(self.vendor_frame, 1, "Address Line 1 (blank = auto)")
        self.vendor_addr2 = self._entry_row(self.vendor_frame, 2, "City, State ZIP (blank = auto)")
        self.vendor_phone = self._entry_row(self.vendor_frame, 3, "Phone (blank = auto)")

        self.options_frame = self._panel("Document Options")
        self.options_frame.grid(row=1, column=0, columnspan=2, padx=20, pady=(0, 20), sticky="ew")
        for i in range(4):
            self.options_frame.grid_columnconfigure(i, weight=1)

        ctk.CTkLabel(self.options_frame, text="Contract Type").grid(row=0, column=0, padx=10, pady=(10, 4), sticky="w")
        self.contract_type_var = ctk.StringVar(value=CONTRACT_TYPES[0])
        ctk.CTkOptionMenu(
            self.options_frame,
            values=CONTRACT_TYPES,
            variable=self.contract_type_var,
            fg_color=self.ACCENT,
            button_color=self.ACCENT,
            button_hover_color=self.ACCENT_HOVER,
            corner_radius=10,
        ).grid(row=1, column=0, padx=10, pady=(0, 10), sticky="ew")

        ctk.CTkLabel(self.options_frame, text="Vertical").grid(row=0, column=1, padx=10, pady=(10, 4), sticky="w")
        self.vertical_var = ctk.StringVar(value=VERTICALS[0])
        ctk.CTkOptionMenu(
            self.options_frame,
            values=VERTICALS,
            variable=self.vertical_var,
            fg_color=self.ACCENT,
            button_color=self.ACCENT,
            button_hover_color=self.ACCENT_HOVER,
            corner_radius=10,
        ).grid(row=1, column=1, padx=10, pady=(0, 10), sticky="ew")

        ctk.CTkLabel(self.options_frame, text="Export Format").grid(row=0, column=2, padx=10, pady=(10, 4), sticky="w")
        self.export_var = ctk.StringVar(value=EXPORT_FORMATS[0])
        ctk.CTkOptionMenu(
            self.options_frame,
            values=EXPORT_FORMATS,
            variable=self.export_var,
            fg_color=self.ACCENT,
            button_color=self.ACCENT,
            button_hover_color=self.ACCENT_HOVER,
            corner_radius=10,
        ).grid(row=1, column=2, padx=10, pady=(0, 10), sticky="ew")

        self.doc_contract = ctk.BooleanVar(value=True)
        self.doc_coi = ctk.BooleanVar(value=True)
        self.doc_amend = ctk.BooleanVar(value=True)

        toggles_frame = ctk.CTkFrame(self.options_frame, fg_color="transparent")
        toggles_frame.grid(row=1, column=3, padx=10, pady=(0, 10), sticky="nsew")
        toggles_frame.grid_columnconfigure(0, weight=1)

        ctk.CTkCheckBox(toggles_frame, text="Create Contract", variable=self.doc_contract).grid(row=0, column=0, sticky="w", pady=2)
        ctk.CTkCheckBox(toggles_frame, text="Create COI", variable=self.doc_coi).grid(row=1, column=0, sticky="w", pady=2)
        ctk.CTkCheckBox(toggles_frame, text="Create Amendment", variable=self.doc_amend).grid(row=2, column=0, sticky="w", pady=2)

        self.exhibits_frame = self._panel("Exhibits (Append to Contract)")
        self.exhibits_frame.grid(row=2, column=0, columnspan=2, padx=20, pady=(0, 20), sticky="ew")
        for i in range(3):
            self.exhibits_frame.grid_columnconfigure(i, weight=1)

        self.exhibit_a = ctk.BooleanVar(value=True)
        self.exhibit_b = ctk.BooleanVar(value=True)
        self.exhibit_c = ctk.BooleanVar(value=True)

        ctk.CTkCheckBox(self.exhibits_frame, text="Exhibit A — Scope", variable=self.exhibit_a).grid(row=0, column=0, padx=10, pady=10, sticky="w")
        ctk.CTkCheckBox(self.exhibits_frame, text="Exhibit B — Pricing", variable=self.exhibit_b).grid(row=0, column=1, padx=10, pady=10, sticky="w")
        ctk.CTkCheckBox(self.exhibits_frame, text="Exhibit C — Security", variable=self.exhibit_c).grid(row=0, column=2, padx=10, pady=10, sticky="w")

        self.output_frame = self._panel("Output Settings")
        self.output_frame.grid(row=3, column=0, columnspan=2, padx=20, pady=(0, 20), sticky="ew")
        self.output_frame.grid_columnconfigure(0, weight=1)
        self.output_frame.grid_columnconfigure(1, weight=0)

        self.out_label = ctk.CTkLabel(self.output_frame, text=f"Output folder:\n{self.out_dir}", justify="left")
        self.out_label.grid(row=0, column=0, padx=10, pady=10, sticky="w")

        ctk.CTkButton(
            self.output_frame,
            text="Choose Export Folder",
            command=self.choose_export_folder,
            fg_color=self.ACCENT,
            hover_color=self.ACCENT_HOVER,
            corner_radius=10,
        ).grid(row=0, column=1, padx=10, pady=10, sticky="e")

        self.disclaimer_label = ctk.CTkLabel(
            self,
            text=APP_DISCLAIMER,
            text_color="#8A5A00",
            fg_color="#FFF4D6",
            corner_radius=10,
            justify="left",
            wraplength=1000,
            padx=12,
            pady=10,
        )
        self.disclaimer_label.grid(row=4, column=0, columnspan=2, padx=20, pady=(0, 12), sticky="ew")

        ctk.CTkButton(
            self,
            text="Generate Test Documents",
            command=self.on_generate,
            fg_color=self.ACCENT,
            hover_color=self.ACCENT_HOVER,
            corner_radius=10,
            height=44,
        ).grid(row=5, column=0, columnspan=2, padx=20, pady=(0, 20), sticky="ew")

        self.status = ctk.CTkLabel(self, text="Ready.", text_color="#666666")
        self.status.grid(row=6, column=0, columnspan=2, padx=20, pady=(0, 10), sticky="w")

    def _panel(self, title: str) -> ctk.CTkFrame:
        frame = ctk.CTkFrame(
            self,
            fg_color=self.PANEL_BG,
            border_color=self.PANEL_BORDER,
            border_width=1,
            corner_radius=10,
        )
        frame.grid_rowconfigure(0, weight=0)
        ctk.CTkLabel(frame, text=title, font=ctk.CTkFont(size=16, weight="bold")).grid(
            row=0, column=0, columnspan=4, padx=16, pady=(12, 8), sticky="w"
        )
        return frame

    def _entry_row(self, parent: ctk.CTkFrame, row_index: int, label: str) -> ctk.CTkEntry:
        ctk.CTkLabel(parent, text=label).grid(row=row_index * 2 + 1, column=0, padx=16, pady=(4, 0), sticky="w")
        entry = ctk.CTkEntry(parent, corner_radius=8)
        entry.grid(row=row_index * 2 + 2, column=0, padx=16, pady=(0, 6), sticky="ew")
        return entry

    def choose_export_folder(self) -> None:
        folder = filedialog.askdirectory()
        if folder:
            self.out_dir = folder
            self.out_label.configure(text=f"Output folder:\n{self.out_dir}")

    def on_generate(self) -> None:
        if not (self.doc_contract.get() or self.doc_coi.get() or self.doc_amend.get()):
            messagebox.showwarning("No documents selected", "Please select at least one document type.")
            return

        def resolve(entry: ctk.CTkEntry, gen_fn) -> str:
            val = entry.get().strip()
            if not val:
                val = gen_fn()
                entry.delete(0, "end")
                entry.insert(0, val)
            return val

        b_addr = rand_address()
        v_addr = rand_address()

        buyer = Party(
            resolve(self.buyer_name, rand_company),
            resolve(self.buyer_addr1, lambda: b_addr[0]),
            resolve(self.buyer_addr2, lambda: b_addr[1]),
            resolve(self.buyer_phone, rand_phone),
        )

        vendor = Party(
            resolve(self.vendor_name, rand_company),
            resolve(self.vendor_addr1, lambda: v_addr[0]),
            resolve(self.vendor_addr2, lambda: v_addr[1]),
            resolve(self.vendor_phone, rand_phone),
        )

        exhibits = [k for k, v in [("A", self.exhibit_a), ("B", self.exhibit_b), ("C", self.exhibit_c)] if v.get()]

        try:
            created = generate_docs(
                self.out_dir,
                buyer,
                vendor,
                self.contract_type_var.get(),
                self.vertical_var.get(),
                self.export_var.get(),
                self.doc_contract.get(),
                self.doc_coi.get(),
                self.doc_amend.get(),
                exhibits,
            )
        except Exception as e:
            messagebox.showerror("Error", str(e))
            return

        messagebox.showinfo("Done", "Created files:\n\n" + "\n".join(created))
        self.status.configure(text=f"Generated {len(created)} file(s). Saved to: {self.out_dir}")


if __name__ == "__main__":
    App().mainloop()
