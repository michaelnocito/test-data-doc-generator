"""DOCX renderer using python-docx.

Layout: title paragraph, party table (2-col borderless), line items table,
totals block, footer paragraph. SAMPLE text in header (no true watermark
support in docx).
"""

import secrets
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from recordforge.core.faker_utils import sanitize_filename
from recordforge.core.models import DocumentData, GeneratedDoc

_FOOTER = (
    "FICTIONAL — For testing, demo, and training use only. "
    "Not a valid legal or financial document."
)


def _remove_table_borders(table) -> None:
    """Set all table borders to none via XML."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)


def render(data: DocumentData, output_dir: Path) -> GeneratedDoc:
    """Render a DocumentData instance to a DOCX file."""
    stem = sanitize_filename(f"{data.doc_type}_{secrets.token_hex(3)}")
    path = output_dir / f"{stem}.docx"

    doc = Document()

    # SAMPLE in header
    section = doc.sections[0]
    hdr_para = section.header.paragraphs[0]
    hdr_run = hdr_para.add_run("SAMPLE")
    hdr_run.font.size = Pt(36)
    hdr_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    hdr_para.alignment = 1  # center

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(data.doc_type.replace("_", " ").title())
    title_run.bold = True
    title_run.font.size = Pt(18)

    # Doc number + date
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(
        f"{data.doc_number}   ·   {data.doc_date}"
        + (f"   ·   Due: {data.due_date}" if data.due_date else "")
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # Party table — 2 columns, borderless
    party_tbl = doc.add_table(rows=1, cols=2)
    _remove_table_borders(party_tbl)

    for cell, label, party in [
        (party_tbl.cell(0, 0), "BILL TO", data.buyer),
        (party_tbl.cell(0, 1), "FROM", data.vendor),
    ]:
        cell.paragraphs[0].clear()
        lbl = cell.paragraphs[0].add_run(label + "\n")
        lbl.font.size = Pt(7)
        lbl.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        lbl.bold = True
        name_run = cell.paragraphs[0].add_run(party.name + "\n")
        name_run.bold = True
        name_run.font.size = Pt(10)
        rest = cell.paragraphs[0].add_run(
            f"{party.address1}\n{party.address2}\n{party.phone}\n{party.email}"
        )
        rest.font.size = Pt(9)

    doc.add_paragraph()

    # Line items table
    if data.line_items:
        headers = ["Description", "Qty", "Unit Price", "Total"]
        n_data = len(data.line_items)
        tbl = doc.add_table(rows=1 + n_data + 3, cols=4)
        tbl.style = "Table Grid"

        # Header row
        for col_idx, hdr in enumerate(headers):
            cell = tbl.cell(0, col_idx)
            cell.text = hdr
            cell.paragraphs[0].runs[0].bold = True

        # Data rows
        for row_idx, item in enumerate(data.line_items, start=1):
            row = tbl.rows[row_idx]
            row.cells[0].text = item.description
            row.cells[1].text = str(item.quantity)
            row.cells[2].text = f"${item.unit_price:,.2f}"
            row.cells[3].text = f"${item.total:,.2f}"

        # Totals rows
        for offset, label, value in [
            (1, "Subtotal", f"${data.subtotal:,.2f}"),
            (2, "Tax (8%)", f"${data.tax:,.2f}"),
            (3, "Total Due", f"${data.total_due:,.2f}"),
        ]:
            row = tbl.rows[n_data + offset]
            row.cells[2].text = label
            row.cells[3].text = value
            if offset == 3:
                for c in (row.cells[2], row.cells[3]):
                    for run in c.paragraphs[0].runs:
                        run.bold = True

        doc.add_paragraph()

    # Notes
    if data.notes:
        doc.add_paragraph(data.notes)
        doc.add_paragraph()

    # Footer
    footer_para = doc.add_paragraph(_FOOTER)
    footer_run = footer_para.runs[0]
    footer_run.font.size = Pt(7)
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    footer_run.font.italic = True

    doc.save(str(path))
    return GeneratedDoc(path=path, doc_type=data.doc_type, format="docx")
